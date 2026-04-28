"""APIs JSON do módulo de Leads (consumidas pelo front via fetch())."""
from __future__ import annotations

import json
import os
import re
import urllib.request
from urllib.parse import quote

from flask import Response, abort, jsonify, render_template, request, send_file, session
from werkzeug.security import check_password_hash

from . import db, leads_api_bp
from .storage import get_storage

# ---------------------------------------------------------------------------
# CNAE baixo risco — carregado uma vez na inicialização
# ---------------------------------------------------------------------------
_BAIXO_RISCO_PATH = os.path.join(
    os.path.dirname(os.path.dirname(__file__)),
    "liberdade economica", "baixo_risco_cnaes.json"
)
_BAIXO_RISCO_SET: set[str] = set()
try:
    with open(_BAIXO_RISCO_PATH, encoding="utf-8") as _f:
        # Normalize: remove hyphen and slash → "6920-6/01" → "6920601"
        _BAIXO_RISCO_SET = {c.replace("-", "").replace("/", "") for c in json.load(_f)}
except Exception:
    pass


def _get_ficha_for_lead(lead_id: str):
    """Load ficha_data for a lead (or its parent organ card)."""
    lead = db.get_lead(lead_id)
    if not lead:
        return None
    if lead.get("ficha_id"):
        from db import get_ficha as _gf
        return _gf(int(lead["ficha_id"]))
    parent = db.get_lead_parent(lead_id)
    if parent and parent.get("ficha_id"):
        from db import get_ficha as _gf
        return _gf(int(parent["ficha_id"]))
    return None


def _tem_atividade_no_local(ficha_data) -> bool:
    """True if the ficha has at least one activity with desenvolvidaNoLocal=True."""
    if not ficha_data:
        return True  # default: assume yes (safer — creates bombeiro)
    dados = ficha_data.get("dados") or {}
    if isinstance(dados, str):
        try:
            dados = __import__("json").loads(dados)
        except Exception:
            return True
    # Check new activities first (alteração with objeto_social change), then empresa_atual, then empresa
    alts = dados.get("alteracoes") or {}
    obj_social = alts.get("objeto_social") or {}
    if obj_social.get("ativo") and obj_social.get("atividades"):
        atividades = obj_social["atividades"]
    else:
        empresa = dados.get("empresa") or {}
        empresa_atual = dados.get("empresa_atual") or {}
        atividades = empresa.get("atividades") or empresa_atual.get("atividades") or []
    return any(bool(a.get("desenvolvidaNoLocal")) for a in atividades)


def _alteration_needs_organ_check(dados: dict) -> bool:
    """True if the alteração changes something other than only nome_empresarial or socios.
    Returns True for constituição (dados has empresa, not empresa_atual).
    Skips only if the ONLY changes are: nome_empresarial, ingresso_socios, retirada_socios,
    transferencia_cotas — i.e. no address, activity, capital, or admin changes."""
    if not dados:
        return True
    # If dados has 'empresa' key without 'empresa_atual', it's a constituição → always check
    if "empresa" in dados and "empresa_atual" not in dados:
        return True
    # Check what alterações are active
    alts = dados.get("alteracoes") or {}
    # These are "organ-triggering" alterations
    organ_triggers = {"endereco", "objeto_social", "capital_social", "administracao", "outras_clausulas"}
    has_organ_trigger = any(
        (alts.get(k) or {}).get("ativo") for k in organ_triggers
    )
    if has_organ_trigger:
        return True
    # Also trigger if there are new socios with capital contribution (ingresso with capital change)
    ingressos = alts.get("ingresso_socios") or []
    if ingressos and (alts.get("capital_social") or {}).get("ativo"):
        return True
    return False


def _analyze_cnaes(ficha_data) -> list[dict]:
    """Returns list of {cnae, descricao, principal, is_baixo_risco} for the ficha.
    Handles both constituição (empresa.atividades) and alteração (empresa_atual.atividades).
    For alterações without CNAE change, falls back to empresa_atual.atividades."""
    if not ficha_data or not _BAIXO_RISCO_SET:
        return []
    dados = ficha_data.get("dados") or {}
    if isinstance(dados, str):
        try:
            dados = json.loads(dados)
        except Exception:
            return []
    # Constituição uses empresa.atividades; alteração uses empresa_atual or empresa
    empresa = dados.get("empresa", {}) or {}
    empresa_atual = dados.get("empresa_atual", {}) or {}
    # Check if this is an alteração with new atividades; if not, use empresa_atual
    atividades = empresa.get("atividades", [])
    if not atividades and empresa_atual:
        atividades = empresa_atual.get("atividades", [])
    result = []
    for ativ in atividades:
        cnae_raw = str(ativ.get("cnae") or "").strip()
        if not cnae_raw:
            continue
        norm = cnae_raw.replace("-", "").replace("/", "").replace(".", "")
        result.append({
            "cnae": cnae_raw,
            "descricao": ativ.get("descricao") or "",
            "principal": bool(ativ.get("principal")),
            "is_baixo_risco": norm in _BAIXO_RISCO_SET,
        })
    return result

# Stage names that trigger guards
STAGE_GUARD_NAMES = {
    "Em Aprovação com Cliente": "client_approval",
    "Conferência Interna": "internal_review",
    "Assinatura do Cliente e Pagamento": "signature",
    "Protocolo na Junta Comercial": "junta",
}

# Stages requiring password for backward movement
PASSWORD_GUARDED_STAGES = {"Assinatura do Cliente e Pagamento", "Em Aprovação com Cliente"}


# ---------------------------------------------------------------------------
# Card modal (partial HTML) + dados auxiliares
# ---------------------------------------------------------------------------

@leads_api_bp.route("/<lead_id>/modal")
def modal_partial(lead_id):
    """Retorna o HTML parcial do conteúdo do modal (lazy load)."""
    lead = db.get_lead(lead_id)
    if not lead:
        abort(404)
    type_data = db.get_lead_type(lead["lead_type_id"])

    ficha_data = None
    parent = db.get_lead_parent(lead_id)

    # For organ cards, also try to load ficha from parent lead
    if lead.get("ficha_id"):
        from db import get_ficha as _get_ficha
        ficha_data = _get_ficha(int(lead["ficha_id"]))
    elif parent and parent.get("ficha_id"):
        from db import get_ficha as _get_ficha
        ficha_data = _get_ficha(int(parent["ficha_id"]))

    cnae_analysis = _analyze_cnaes(ficha_data)
    any_baixo_risco = any(a["is_baixo_risco"] for a in cnae_analysis)
    all_baixo_risco = bool(cnae_analysis) and all(a["is_baixo_risco"] for a in cnae_analysis)

    # Current stage for SLA display
    stages = db.list_stages(lead["workflow_id"])
    current_stage_obj = next((s for s in stages if s["id"] == lead.get("current_stage_id")), None)

    from db import list_users as _list_users
    # Stage checklist templates for the current stage (to show unapplied suggestions)
    stage_templates = []
    if lead.get("current_stage_id"):
        stage_templates = db.list_stage_checklist_templates(lead["current_stage_id"])
        # Also include items from linked checklist_template_id
        with db.db_cursor() as _conn:
            row = _conn.execute(
                "SELECT checklist_template_id FROM lead_stages WHERE id=?",
                (lead["current_stage_id"],)
            ).fetchone()
        if row and row[0]:
            tpl = db.get_checklist_template(row[0])
            if tpl:
                for item in tpl.get("items", []):
                    stage_templates.append({"label": item["label"], "required": item.get("required", 0)})
    return render_template(
        "leads/_card_modal.html",
        lead=lead,
        type_data=type_data,
        stages=stages,
        macrophases=db.list_macrophases(lead["workflow_id"]),
        priorities=db.list_priorities(),
        statuses=db.list_statuses(),
        tags=db.list_tags(),
        offices=db.list_offices(),
        ficha_data=ficha_data,
        comments=db.list_comments(lead_id),
        history=db.list_history(lead_id),
        files=db.list_files(lead_id),
        checklist=db.list_checklist(lead_id),
        stage_templates=stage_templates,
        form_fields=db.get_form_fields(lead["lead_type_id"]),
        users=_list_users(),
        children=db.get_lead_children(lead_id),
        parent=parent,
        cnae_analysis=cnae_analysis,
        any_baixo_risco=any_baixo_risco,
        all_baixo_risco=all_baixo_risco,
        current_stage_obj=current_stage_obj,
    )


# ---------------------------------------------------------------------------
# Atualização de campos do lead (Operacional + Kanban DnD)
# ---------------------------------------------------------------------------

@leads_api_bp.route("/<lead_id>", methods=["PATCH"])
def patch_lead(lead_id):
    payload = request.get_json(silent=True) or {}
    actor = session.get("user_name") or "Sistema"
    db.update_lead_fields(lead_id, payload, actor=actor)
    if "tag_ids" in payload:
        db.set_lead_tags(lead_id, payload["tag_ids"] or [])
    return jsonify(db.get_lead(lead_id))


@leads_api_bp.route("/<lead_id>/move", methods=["POST"])
def move_card(lead_id):
    data = request.get_json(silent=True) or {}
    stage_id = data.get("stage_id")
    if not stage_id:
        return jsonify({"error": "stage_id obrigatório"}), 400

    lead = db.get_lead(lead_id)
    if not lead:
        abort(404)

    actor = session.get("user_name") or "Sistema"

    current_stage = db.get_stage(lead["current_stage_id"]) if lead.get("current_stage_id") else None
    target_stage  = db.get_stage(stage_id)

    if current_stage and target_stage:
        current_pos = current_stage.get("position", 0)
        target_pos  = target_stage.get("position", 0)
        moving_forward  = target_pos > current_pos
        moving_backward = target_pos < current_pos

        current_stage_name = current_stage.get("name", "")
        target_stage_name  = target_stage.get("name", "")

        justification  = (data.get("justification") or "").strip()
        guard_password = (data.get("guard_password") or "").strip()

        # Helper: validate gerente/admin password
        def _check_guard_password(pwd: str) -> bool:
            if not pwd:
                return False
            profile = session.get("profile", "")
            if profile not in ("admin", "gerente"):
                return False
            user_id = session.get("user_id")
            if not user_id:
                return False
            try:
                from db import get_user_by_id as _get_user
                user = _get_user(user_id)
                if user and check_password_hash(user["password_hash"], pwd):
                    return True
            except Exception:
                pass
            return False

        # --- Guard: Moving BACKWARD from a guarded stage ---
        if moving_backward and current_stage_name in STAGE_GUARD_NAMES:
            guard_type = STAGE_GUARD_NAMES[current_stage_name]

            # Password check for specific stages
            if current_stage_name in PASSWORD_GUARDED_STAGES:
                if not guard_password:
                    return jsonify({
                        "require": "justification",
                        "guard_type": guard_type,
                        "needs_password": True,
                        "message": f"Retroceder de '{current_stage_name}' requer justificativa e senha de gerente/admin.",
                    }), 409
                if not _check_guard_password(guard_password):
                    return jsonify({
                        "require": "justification",
                        "guard_type": guard_type,
                        "needs_password": True,
                        "message": "Senha de gerente/admin incorreta.",
                    }), 409

            if not justification:
                return jsonify({
                    "require": "justification",
                    "guard_type": guard_type,
                    "needs_password": current_stage_name in PASSWORD_GUARDED_STAGES,
                    "message": f"É necessária uma justificativa para retroceder de '{current_stage_name}'.",
                }), 409

            # Log guard event
            db.log_guard_event(
                lead_id=lead_id,
                event_type=f"backward_{guard_type}",
                stage_from=lead["current_stage_id"],
                stage_to=stage_id,
                actor_name=actor,
                justification=justification,
            )

        # --- Guard: Moving FORWARD from "Em Aprovação com Cliente" ---
        if moving_forward and current_stage_name == "Em Aprovação com Cliente":
            # Check for approved client_approval OR gerente/admin override
            approval = db.get_lead_approval(lead_id, "client_approval")
            has_approved = approval and approval.get("status") == "approved"
            has_override = _check_guard_password(guard_password)

            if not has_approved and not has_override:
                return jsonify({
                    "require": "client_approval",
                    "message": "Aguardando aprovação do cliente.",
                }), 409

        # --- Guard: Required checklist items must be completed before advancing ---
        if moving_forward and lead.get("current_stage_id"):
            chk_ok, missing_items = db.check_stage_checklist_complete(
                lead_id, lead["current_stage_id"])
            if not chk_ok:
                preview = missing_items[:3]
                suffix = f" (e mais {len(missing_items)-3})" if len(missing_items) > 3 else ""
                return jsonify({
                    "require": "checklist",
                    "message": (
                        f"Conclua os itens obrigatórios do checklist antes de avançar: "
                        f"{', '.join(preview)}{suffix}."
                    ),
                    "missing": missing_items,
                }), 409

        # --- Guard: Organ child card leaving "Protocolo do Pedido" — requires protocol data ---
        if (moving_forward and lead.get("organ_type")
                and current_stage and "protocolo do pedido" in current_stage_name.lower()):
            organ_type = lead["organ_type"]
            try:
                organs_raw = lead.get("op_organs_data") or "{}"
                organs_data = json.loads(organs_raw)
            except Exception:
                organs_data = {}
            organ_field = f"op_{organ_type}"
            organ_info  = organs_data.get(organ_field, {})
            if not organ_info.get("protocolo", "").strip():
                label_map = {
                    "bombeiro":   "Bombeiro",
                    "vigilancia": "Vigilância Sanitária",
                    "conselho":   "Conselho de Classe",
                    "alvara":     "Alvará",
                }
                label = label_map.get(organ_type, organ_type.title())
                return jsonify({
                    "require": "protocol_data",
                    "message": (
                        f"Preencha o número de protocolo na aba Órgãos ({label}) "
                        f"antes de avançar da etapa Protocolo do Pedido."
                    ),
                }), 409

    db.update_lead_fields(lead_id, {"current_stage_id": stage_id}, actor=actor)
    # Auto-populate checklist items from stage templates
    db.apply_stage_checklist_templates(lead_id, stage_id)
    lead = db.get_lead(lead_id)

    # Determine if junta organ modal should appear (leaving "Protocolo na Junta Comercial" forward)
    show_junta_modal = False
    _moving_forward = (
        current_stage and target_stage and
        target_stage.get("position", 0) > current_stage.get("position", 0)
    )
    if (current_stage and current_stage.get("name") == "Protocolo na Junta Comercial"
            and _moving_forward):
        ficha_data = _get_ficha_for_lead(lead_id)
        dados = {}
        if ficha_data:
            d = ficha_data.get("dados") or {}
            dados = json.loads(d) if isinstance(d, str) else d
        show_junta_modal = _alteration_needs_organ_check(dados)

    return jsonify({"ok": True, "stage_entered_at": lead["stage_entered_at"],
                    "show_junta_modal": show_junta_modal})


@leads_api_bp.route("/<lead_id>/generate-approval", methods=["POST"])
def generate_approval(lead_id):
    """Creates or returns existing pending approval for a lead."""
    lead = db.get_lead(lead_id)
    if not lead:
        abort(404)

    # Check for existing pending approval
    existing = db.get_lead_approval(lead_id, "client_approval")
    if existing and existing.get("status") == "pending":
        token = existing["token"]
        access_code = existing.get("access_code", "1234")
    else:
        # Create new approval
        import uuid as _uuid
        token = _uuid.uuid4().hex
        db.create_approval(lead_id, "client_approval", token=token, access_code="1234")
        access_code = "1234"

    base_url = request.host_url.rstrip("/")
    link = f"{base_url}/leads/aprovacao/{token}"
    return jsonify({"ok": True, "token": token, "link": link, "access_code": access_code})


@leads_api_bp.route("/approval/<token>/resolve", methods=["POST"])
def resolve_approval_route(token):
    """Public endpoint — resolves a client approval."""
    approval = db.get_approval_by_token(token)
    if not approval:
        abort(404)
    if approval.get("status") != "pending":
        return jsonify({"error": "Esta aprovação já foi resolvida."}), 400

    data = request.get_json(silent=True) or {}
    action          = data.get("action")  # 'approve' or 'reject'
    access_code     = (data.get("access_code") or "").strip()
    justification   = (data.get("justification") or "").strip()
    return_stage_id = data.get("return_stage_id") or None
    client_name     = (data.get("client_name") or "").strip()
    client_cpf      = (data.get("client_cpf") or "").strip()

    if access_code != (approval.get("access_code") or "1234"):
        return jsonify({"error": "Código de acesso incorreto."}), 403

    if action not in ("approve", "reject"):
        return jsonify({"error": "Ação inválida."}), 400

    lead_id = approval["lead_id"]
    lead = db.get_lead(lead_id)
    if not lead:
        abort(404)

    # Build actor label with name/CPF if provided
    actor_label = "cliente"
    if client_name:
        actor_label = f"{client_name} (CPF: {client_cpf})" if client_cpf else client_name

    if action == "approve":
        full_justification = justification
        if client_name or client_cpf:
            id_line = f"Aprovado por: {actor_label}"
            full_justification = f"{id_line}\n{justification}".strip()
        db.resolve_approval(approval["id"], "approved", justification=full_justification,
                            resolved_by=actor_label)
        # Record in lead history as a comment
        hist_body = f"✅ Aprovação do cliente\nNome: {client_name or '—'}\nCPF: {client_cpf or '—'}"
        if justification:
            hist_body += f"\nObservação: {justification}"
        db.add_comment(lead_id, hist_body, author=actor_label)
        # Advance lead to next stage
        stages = db.list_stages(lead["workflow_id"])
        current_pos = None
        for st in stages:
            if st["id"] == lead["current_stage_id"]:
                current_pos = st["position"]
                break
        next_stage = None
        if current_pos is not None:
            for st in stages:
                if st["position"] > current_pos:
                    next_stage = st
                    break
        if next_stage:
            db.update_lead_fields(lead_id, {"current_stage_id": next_stage["id"]})
            db.log_guard_event(lead_id, "client_approved",
                               stage_from=lead["current_stage_id"],
                               stage_to=next_stage["id"],
                               actor_name=actor_label,
                               justification=full_justification or None)
        # Notify responsible
        if lead.get("responsible_name"):
            _notify_by_name(lead["responsible_name"], lead_id, "approval",
                            f'Cliente APROVOU o processo "{lead["name"]}" ({actor_label})',
                            actor_name=actor_label)
        return jsonify({"ok": True, "redirect_stage": next_stage["id"] if next_stage else None})

    else:  # reject
        # Auto-find the previous stage if return_stage_id not provided
        if not return_stage_id:
            stages = db.list_stages(lead["workflow_id"])
            current_pos = None
            for st in stages:
                if st["id"] == lead["current_stage_id"]:
                    current_pos = st["position"]
                    break
            prev_stage = None
            if current_pos is not None:
                for st in reversed(stages):
                    if st["position"] < current_pos:
                        prev_stage = st
                        break
            if prev_stage:
                return_stage_id = prev_stage["id"]
            else:
                return_stage_id = lead["current_stage_id"]  # stay in place if no previous
        full_justification = justification
        if client_name or client_cpf:
            id_line = f"Reprovado por: {actor_label}"
            full_justification = f"{id_line}\n{justification}".strip()
        db.resolve_approval(approval["id"], "rejected", justification=full_justification,
                            return_stage_id=return_stage_id, resolved_by=actor_label)
        # Record in lead history as a comment
        hist_body = f"Reprovação do cliente\nNome: {client_name or '—'}\nCPF: {client_cpf or '—'}\nMotivo: {justification}"
        db.add_comment(lead_id, hist_body, author=actor_label)
        db.update_lead_fields(lead_id, {"current_stage_id": return_stage_id})
        db.log_guard_event(lead_id, "client_rejected",
                           stage_from=lead["current_stage_id"],
                           stage_to=return_stage_id,
                           actor_name=actor_label,
                           justification=full_justification or None)
        # Notify responsible
        if lead.get("responsible_name"):
            _notify_by_name(lead["responsible_name"], lead_id, "rejection",
                            f'Cliente REPROVOU o processo "{lead["name"]}" ({actor_label})',
                            actor_name=actor_label)
        return jsonify({"ok": True, "redirect_stage": return_stage_id})


def _gerar_docx_buf(ficha):
    """Generate DOCX bytes for a ficha (constituicao or alteracao)."""
    import io
    dados = ficha.get("dados") or {}
    tipo = ficha.get("tipo", "constituicao")
    buf = io.BytesIO()
    if tipo == "constituicao":
        from gerar_contrato import gerar_contrato
        gerar_contrato(dados, buf)
    else:
        from gerar_alteracao import gerar_alteracao
        gerar_alteracao(dados, buf)
    buf.seek(0)
    return buf, dados, tipo


def _docx_to_html_page(buf, titulo: str) -> str:
    """Convert DOCX buffer to a faithful HTML page using python-docx (preserves formatting)."""
    from docx import Document
    from docx.text.paragraph import Paragraph as _Para
    from docx.table import Table as _Table
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import html as _html

    doc = Document(buf)
    section = doc.sections[0]

    # Page margins in cm
    lm = (section.left_margin.cm  if section.left_margin  else 1.99)
    rm = (section.right_margin.cm if section.right_margin else 1.75)
    tm = (section.top_margin.cm   if section.top_margin   else 2.88)
    bm = (section.bottom_margin.cm if section.bottom_margin else 2.50)

    EMU_PER_CM = 914400 / 2.54
    EMU_PER_PT = 12700

    def emu_to_cm(emu):
        return emu / EMU_PER_CM if emu else 0

    def emu_to_pt(emu):
        return emu / EMU_PER_PT if emu else 0

    def render_run(run):
        text = _html.escape(run.text)
        if not text:
            return ""
        # Font size
        fs = run.font.size
        size_pt = emu_to_pt(fs) if fs else None
        style = f"font-size:{size_pt:.1f}pt;" if size_pt and abs(size_pt - 12) > 0.5 else ""
        # Color
        color = run.font.color.rgb if (run.font.color and run.font.color.type) else None
        if color:
            style += f"color:#{color};"
        if style:
            text = f'<span style="{style}">{text}</span>'
        if run.bold:
            text = f"<strong>{text}</strong>"
        if run.italic:
            text = f"<em>{text}</em>"
        if run.underline:
            text = f"<u>{text}</u>"
        return text

    def render_para(para):
        # Alignment
        align_map = {
            WD_ALIGN_PARAGRAPH.CENTER:  "center",
            WD_ALIGN_PARAGRAPH.RIGHT:   "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
            WD_ALIGN_PARAGRAPH.LEFT:    "left",
        }
        align = align_map.get(para.alignment, "justify")

        pf = para.paragraph_format
        left_cm  = emu_to_cm(pf.left_indent)  if pf.left_indent  else 0
        first_cm = emu_to_cm(pf.first_line_indent) if pf.first_line_indent else 0
        sa_pt    = emu_to_pt(pf.space_after)   if pf.space_after  else 4
        sb_pt    = emu_to_pt(pf.space_before)  if pf.space_before else 0

        content = "".join(render_run(r) for r in para.runs)
        if not content.strip():
            sp = max(sa_pt, 2)
            return f'<p style="margin:0 0 {sp:.0f}pt 0;">&nbsp;</p>'

        style = f"text-align:{align};margin:0 0 {sa_pt:.0f}pt 0;"
        if sb_pt > 0:
            style += f"margin-top:{sb_pt:.0f}pt;"
        if left_cm > 0:
            style += f"margin-left:{left_cm:.3f}cm;"
        if first_cm != 0:
            style += f"text-indent:{first_cm:.3f}cm;"

        return f'<p style="{style}">{content}</p>\n'

    def render_table(tbl):
        parts = ['<table>']
        for row in tbl.rows:
            parts.append('<tr>')
            for cell in row.cells:
                cell_html = ''
                for child in cell._tc:
                    tag = child.tag.split('}')[-1]
                    if tag == 'p':
                        cell_html += render_para(_Para(child, cell))
                    elif tag == 'tbl':
                        cell_html += render_table(_Table(child, cell))
                parts.append(f'<td>{cell_html}</td>')
            parts.append('</tr>')
        parts.append('</table>\n')
        return ''.join(parts)

    # Iterate body elements in order
    body_html = []
    for child in doc.element.body:
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            body_html.append(render_para(_Para(child, doc)))
        elif tag == 'tbl':
            body_html.append(render_table(_Table(child, doc)))

    body_str = ''.join(body_html)
    titulo_esc = _html.escape(titulo)

    return f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>{titulo_esc}</title>
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{background:#d0d0d0;padding:20px;font-family:Arial,Helvetica,sans-serif;font-size:12pt}}
.doc-page{{
  background:#fff;
  width:21cm;
  min-height:29.7cm;
  margin:0 auto;
  padding:{tm:.2f}cm {rm:.2f}cm {bm:.2f}cm {lm:.2f}cm;
  box-shadow:0 4px 24px rgba(0,0,0,.3);
  line-height:1.5;
  word-break:break-word;
}}
.doc-page p{{font-family:Arial,Helvetica,sans-serif;font-size:12pt;}}
.doc-page strong{{font-weight:700;}}
.doc-page table{{width:100%;border-collapse:collapse;margin:6pt 0;font-size:11pt;}}
.doc-page td,.doc-page th{{border:1px solid #555;padding:3pt 5pt;vertical-align:top;}}
.print-btn{{
  position:fixed;bottom:24px;right:24px;
  background:#2456a4;color:#fff;border:none;
  padding:10px 22px;border-radius:8px;
  cursor:pointer;font-size:11pt;
  box-shadow:0 4px 14px rgba(0,0,0,.3);z-index:100;
  font-family:Arial,sans-serif;
}}
.print-btn:hover{{background:#1a3f84}}
@media print{{
  body{{background:#fff;padding:0}}
  .doc-page{{width:100%;box-shadow:none;padding:{tm:.2f}cm {rm:.2f}cm {bm:.2f}cm {lm:.2f}cm;min-height:0}}
  .print-btn{{display:none!important}}
}}
@media(max-width:680px){{
  body{{padding:4px}}
  .doc-page{{width:100%;padding:12px 10px}}
}}
</style>
</head>
<body>
<button class="print-btn" onclick="window.print()">&#128438; Imprimir / PDF</button>
<div class="doc-page">{body_str}</div>
</body>
</html>"""


@leads_api_bp.route("/approval/<token>/documento")
def approval_documento(token):
    """Public endpoint — view the contract/alteration document for an approval link."""
    approval = db.get_approval_by_token(token)
    if not approval:
        abort(404)
    lead = db.get_lead(approval["lead_id"])
    if not lead or not lead.get("ficha_id"):
        abort(404)
    from db import get_ficha as _get_ficha
    ficha = _get_ficha(int(lead["ficha_id"]))
    if not ficha:
        abort(404)
    buf, dados, tipo = _gerar_docx_buf(ficha)
    if tipo == "constituicao":
        titulo = dados.get("empresa", {}).get("razaoSocial", "Contrato")
    else:
        titulo = dados.get("empresa_atual", {}).get("razaoSocial", "Alteração")
    return _docx_to_html_page(buf, titulo), 200, {"Content-Type": "text/html; charset=utf-8"}


@leads_api_bp.route("/<lead_id>/contrato-html")
def contrato_html(lead_id):
    """Logged-in endpoint — view contract as HTML for a lead."""
    lead = db.get_lead(lead_id)
    if not lead or not lead.get("ficha_id"):
        abort(404)
    from db import get_ficha as _get_ficha
    ficha = _get_ficha(int(lead["ficha_id"]))
    if not ficha:
        abort(404)
    buf, dados, tipo = _gerar_docx_buf(ficha)
    if tipo == "constituicao":
        titulo = dados.get("empresa", {}).get("razaoSocial", "Contrato")
    else:
        titulo = dados.get("empresa_atual", {}).get("razaoSocial", "Alteração")
    return _docx_to_html_page(buf, titulo), 200, {"Content-Type": "text/html; charset=utf-8"}


def _do_create_organ_child(lead, organ_key: str) -> dict | None:
    """Create one organ child card for the given organ_key. Returns {id, organ} or None if already exists."""
    organ_labels = {
        "bombeiro":   "Bombeiro",
        "vigilancia": "Vigilância Sanitária",
        "alvara":     "Alvará",
        "conselho":   "Conselho de Classe",
    }
    organ_name = organ_labels.get(organ_key, organ_key.title())
    # Don't duplicate: check if a child with this organ_type already exists
    existing = db.get_lead_children(lead["id"])
    for child in existing:
        if child.get("organ_type") == organ_key:
            return None  # already exists
    organ_type_rec = db.get_lead_type_by_code(organ_key)
    organ_lead_type_id = organ_type_rec["id"] if organ_type_rec else lead["lead_type_id"]
    new_lead_id = db.create_organ_lead(
        parent_lead_id=lead["id"],
        organ_type=organ_key,
        name=f"{lead['name']} — {organ_name}",
        lead_type_id=organ_lead_type_id,
        responsible_name=lead.get("responsible_name"),
        office_id=lead.get("office_id"),
    )
    return {"id": new_lead_id, "organ": organ_name}


@leads_api_bp.route("/<lead_id>/create-organ-leads", methods=["POST"])
def create_organ_leads(lead_id):
    data = request.get_json(silent=True) or {}
    dispensa = data.get("dispensa_licencas", False)
    lead = db.get_lead(lead_id)
    if not lead:
        abort(404)

    # Load ficha to check atividade no local
    ficha_data = _get_ficha_for_lead(lead_id)
    tem_local = _tem_atividade_no_local(ficha_data)

    if dispensa:
        # Junta emitiu certidão de dispensa → only alvará (if flag) + conselho (always)
        db.update_lead_fields(lead_id, {"op_baixo_risco": "sim"})
        organs_to_create = []
        if lead.get("op_alvara") == "sim":
            organs_to_create.append("alvara")
        organs_to_create.append("conselho")   # conselho is always required
    else:
        # Junta NÃO emitiu certidão → force bombeiro (if atividade no local), vigilância, conselho
        organs_to_create = []
        if tem_local:
            organs_to_create.append("bombeiro")   # force, regardless of flag
        organs_to_create.append("vigilancia")     # force, regardless of flag / baixo risco
        organs_to_create.append("conselho")       # always required
        if lead.get("op_alvara") == "sim":
            organs_to_create.append("alvara")

    created = []
    for organ_key in organs_to_create:
        result = _do_create_organ_child(lead, organ_key)
        if result:
            created.append(result)

    return jsonify({"ok": True, "created": created})


@leads_api_bp.route("/<lead_id>/create-organ-child", methods=["POST"])
def create_organ_child_manual(lead_id):
    """Manually create a single organ child card (bombeiro / vigilancia / conselho / alvara)."""
    lead = db.get_lead(lead_id)
    if not lead:
        abort(404)
    data = request.get_json(silent=True) or {}
    organ_key = (data.get("organ_type") or "").strip().lower()
    valid_organs = {"bombeiro", "vigilancia", "conselho", "alvara"}
    if organ_key not in valid_organs:
        return jsonify({"error": "organ_type inválido"}), 400
    result = _do_create_organ_child(lead, organ_key)
    if result is None:
        return jsonify({"ok": True, "created": None, "message": "Card já existe."})
    return jsonify({"ok": True, "created": result})


@leads_api_bp.route("/<lead_id>", methods=["DELETE"])
def delete_lead_route(lead_id):
    db.delete_lead(lead_id)
    return jsonify({"ok": True})


@leads_api_bp.route("/<lead_id>/apply-stage-checklist", methods=["POST"])
def apply_stage_checklist(lead_id):
    """Apply the current stage's checklist templates to this lead (for leads already in the stage)."""
    lead = db.get_lead(lead_id)
    if not lead:
        abort(404)
    stage_id = lead.get("current_stage_id")
    if not stage_id:
        return jsonify({"ok": True, "applied": 0})
    before = len(db.list_checklist(lead_id))
    db.apply_stage_checklist_templates(lead_id, stage_id)
    after = len(db.list_checklist(lead_id))
    items = db.list_checklist(lead_id)
    return jsonify({"ok": True, "applied": after - before, "items": items})


@leads_api_bp.route("/manuais-estado/<state_code>")
def download_state_manual(state_code):
    """Public endpoint to download a state signature manual (no auth required — for client portal)."""
    rec = db.get_state_manual(state_code.upper())
    if not rec:
        abort(404)
    stream = get_storage().open_stream(rec["storage_key"])
    import mimetypes as _mt
    mime = _mt.guess_type(rec["filename"])[0] or "application/octet-stream"
    return Response(
        stream,
        mimetype=mime,
        headers={"Content-Disposition": f'attachment; filename="{rec["filename"]}"'},
    )


@leads_api_bp.route("/<lead_id>/children")
def lead_children(lead_id):
    """Returns child organ leads for a parent process."""
    return jsonify(db.get_lead_children(lead_id))


@leads_api_bp.route("/<lead_id>/change-status", methods=["POST"])
def change_status(lead_id):
    """Change lead status.
    - Closed statuses (Cancelado/Inativo): require comment + auto-move to last stage.
    - Pause statuses (Aguardando Cliente/Órgão): require comment, no stage change.
    """
    lead = db.get_lead(lead_id)
    if not lead:
        abort(404)
    data = request.get_json(silent=True) or {}
    new_status = (data.get("status") or "").strip()
    comment_body = (data.get("comment") or "").strip()
    actor = session.get("user_name") or "Sistema"

    if not new_status:
        return jsonify({"error": "Status inválido."}), 400

    # All statuses in COMMENT_REQUIRED_STATUS_NAMES require a comment
    if new_status in db.COMMENT_REQUIRED_STATUS_NAMES and not comment_body:
        return jsonify({"require": "comment", "message": "Informe um motivo para este status."}), 409

    # Save comment if provided
    if comment_body:
        db.add_comment(lead_id, f"[Status: {new_status}] {comment_body}", author=actor)

    # Auto-move to last stage only for definitively closed statuses
    updates: dict = {"status": new_status}
    if new_status in db.CLOSED_STATUS_NAMES:
        last_stage = db.get_last_stage(lead["workflow_id"])
        if last_stage and last_stage["id"] != lead.get("current_stage_id"):
            updates["current_stage_id"] = last_stage["id"]

    db.update_lead_fields(lead_id, updates, actor=actor)

    # Notify responsible user for closed/pause statuses
    if new_status in db.COMMENT_REQUIRED_STATUS_NAMES and lead.get("responsible_name"):
        _notify_by_name(lead.get("responsible_name"), lead_id,
                        "status_change",
                        f'Processo "{lead["name"]}" marcado como {new_status}',
                        actor_name=actor)

    return jsonify({"ok": True, "status": new_status})


# ---------------------------------------------------------------------------
# Formulário (ficha)
# ---------------------------------------------------------------------------

@leads_api_bp.route("/<lead_id>/form", methods=["PUT"])
def save_form(lead_id):
    data = request.get_json(silent=True) or {}
    db.save_form_data(lead_id, data)
    return jsonify({"ok": True})


# ---------------------------------------------------------------------------
# Comentários, histórico, checklist
# ---------------------------------------------------------------------------

@leads_api_bp.route("/<lead_id>/comments", methods=["POST"])
def add_comment(lead_id):
    # Aceita tanto JSON (texto simples) quanto multipart/form-data (com anexo)
    session_author = session.get("user_name") or None
    if request.content_type and 'multipart' in request.content_type:
        body = (request.form.get("body") or "").strip()
        author = request.form.get("author") or session_author
        att_key = att_name = att_mime = None
        if "attachment" in request.files:
            f = request.files["attachment"]
            if f and f.filename:
                storage = get_storage()
                att_key, _ = storage.save(lead_id, f"comment_{f.filename}", f.stream, f.mimetype)
                att_name = f.filename
                att_mime = f.mimetype
    else:
        data = request.get_json(silent=True) or {}
        body = (data.get("body") or "").strip()
        author = data.get("author") or session_author
        att_key = att_name = att_mime = None

    if not body:
        return jsonify({"error": "Comentário vazio."}), 400
    cid = db.add_comment(lead_id, body, author=author,
                         attachment_key=att_key, attachment_name=att_name,
                         attachment_mime=att_mime)

    # Parse @mentions and notify
    lead = db.get_lead(lead_id)
    if lead:
        mentions = re.findall(r'@([\w\s]+?)(?=\s|$|[,;!?])', body)
        users = db.list_users()
        name_to_user = {u["name"].lower(): u for u in users}
        notified = set()
        for mention in mentions:
            key = mention.strip().lower()
            user = name_to_user.get(key)
            if user and user["id"] not in notified:
                notified.add(user["id"])
                db.create_notification(
                    user_id=user["id"],
                    lead_id=lead_id,
                    notif_type="mention",
                    message=f'{author or "Alguém"} mencionou você no processo "{lead["name"]}"',
                    actor_name=author,
                )

    return jsonify({"id": cid, "attachment_name": att_name, "attachment_key": att_key,
                    "attachment_mime": att_mime})


@leads_api_bp.route("/<lead_id>/checklist", methods=["GET", "POST"])
def checklist(lead_id):
    if request.method == "GET":
        return jsonify(db.list_checklist(lead_id))
    data = request.get_json(silent=True) or {}
    label = (data.get("label") or "").strip()
    if not label:
        return jsonify({"error": "Label vazio."}), 400
    required = bool(data.get("required", False))
    cid = db.add_checklist_item(lead_id, label, stage_id=data.get("stage_id"), required=required)
    return jsonify({"id": cid})


@leads_api_bp.route("/checklist/<item_id>", methods=["PATCH", "DELETE"])
def checklist_item(item_id):
    if request.method == "DELETE":
        db.delete_checklist_item(item_id)
        return jsonify({"ok": True})
    data = request.get_json(silent=True) or {}
    db.toggle_checklist_item(item_id, bool(data.get("done")))
    return jsonify({"ok": True})


# ---------------------------------------------------------------------------
# Arquivos (upload via storage driver, download por streaming)
# ---------------------------------------------------------------------------

@leads_api_bp.route("/<lead_id>/files", methods=["POST"])
def upload_file(lead_id):
    if "file" not in request.files:
        return jsonify({"error": "Sem arquivo."}), 400
    f = request.files["file"]
    storage = get_storage()
    storage_key, size = storage.save(
        lead_id, f.filename or "arquivo", f.stream, f.mimetype
    )
    fid = db.add_file(
        lead_id,
        filename=f.filename or "arquivo",
        storage_key=storage_key,
        size_bytes=size,
        mime_type=f.mimetype,
    )
    return jsonify({"id": fid, "filename": f.filename, "size_bytes": size})


@leads_api_bp.route("/comment-attachment/<path:storage_key>")
def download_comment_attachment(storage_key):
    stream = get_storage().open_stream(storage_key)
    filename = storage_key.split("/")[-1].split("comment_", 1)[-1]
    return Response(stream, mimetype="application/octet-stream",
                    headers={"Content-Disposition": f'attachment; filename="{filename}"'})


@leads_api_bp.route("/files/<file_id>")
def download_file(file_id):
    rec = db.get_file(file_id)
    if not rec:
        abort(404)
    stream = get_storage().open_stream(rec["storage_key"])
    return Response(
        stream,
        mimetype=rec["mime_type"] or "application/octet-stream",
        headers={
            "Content-Disposition":
                f'attachment; filename="{rec["filename"]}"',
        },
    )


@leads_api_bp.route("/files/<file_id>", methods=["DELETE"])
def remove_file(file_id):
    rec = db.delete_file(file_id)
    if rec:
        try:
            get_storage().delete(rec["storage_key"])
        except Exception:
            pass  # idempotente — registro já foi removido
    return jsonify({"ok": True})


# ---------------------------------------------------------------------------
# Geração de declarações e documentos por IA
# ---------------------------------------------------------------------------

def _build_decl_data(lead, ficha_data, parent):
    """Extrai dados do formulário (constituição ou alteração) para as declarações."""
    from datetime import date as _date

    fd = ficha_data or {}
    dados_form = fd.get("dados") or {}
    # Alteração usa empresa_atual; constituição usa empresa
    empresa = (
        dados_form.get("empresa_atual")
        or dados_form.get("empresa")
        or {}
    )
    end = empresa.get("enderecoComercial") or {}

    log_tipo = (end.get("logradouroTipo") or "").strip()
    log_desc = (end.get("logradouroDescricao") or "").strip()
    logradouro = f"{log_tipo} {log_desc}".strip() or end.get("logradouro") or ""

    socios = empresa.get("socios") or []
    admin  = next((s for s in socios if s.get("administrador")), socios[0] if socios else {})

    meses = ["janeiro","fevereiro","março","abril","maio","junho",
             "julho","agosto","setembro","outubro","novembro","dezembro"]
    hoje = _date.today()
    data_hoje = f"{hoje.day:02d} de {meses[hoje.month - 1]} de {hoje.year}"

    return {
        "razao_social":  empresa.get("razaoSocial") or lead.get("name") or "",
        "cnpj":          empresa.get("cnpj") or "[CNPJ a definir após registro]",
        "cidade":        end.get("cidade") or "",
        "estado":        end.get("estado") or "PR",
        "logradouro":    logradouro,
        "numero":        end.get("numero") or "s/nº",
        "complemento":   end.get("complemento") or "",
        "bairro":        end.get("bairro") or "",
        "cep":           end.get("cep") or "",
        "admin_nome":    admin.get("nome") or "",
        "admin_cpf":     admin.get("cpf") or "",
        "admin_cargo":   "Sócio(a) Administrador(a)",
        "data_hoje":     data_hoje,
    }


def _make_docx_declaration(tipo: str, d: dict) -> bytes:
    """Gera .docx em padrão ABNT — Arial 12, justificado, espaçamento 1,5."""
    import io
    import re as _re
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    FONT   = "Arial"
    SIZE   = Pt(12)
    SPACE  = Pt(18)   # 1,5 × 12 pt
    INDENT = Cm(1.25)

    doc = Document()

    # Margens ABNT
    for sec in doc.sections:
        sec.top_margin    = Cm(3)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(3)
        sec.right_margin  = Cm(2)

    def _apply_font(run):
        """Arial 12 via API + XML (cobre scripts complexos)."""
        run.font.name = FONT
        run.font.size = SIZE
        rPr = run._r.get_or_add_rPr()
        # Remove any existing rFonts before inserting ours
        for old in rPr.findall(qn('w:rFonts')):
            rPr.remove(old)
        rf = OxmlElement('w:rFonts')   # prefixed tag, NOT qn() here
        rf.set(qn('w:ascii'), FONT)
        rf.set(qn('w:hAnsi'), FONT)
        rf.set(qn('w:cs'),    FONT)
        rPr.insert(0, rf)

    def para(text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             sb=0, sa=0, indent=None):
        p = doc.add_paragraph()
        p.alignment = align
        fmt = p.paragraph_format
        fmt.line_spacing  = SPACE
        fmt.space_before  = Pt(sb)
        fmt.space_after   = Pt(sa)
        if indent is not None:
            fmt.first_line_indent = indent
        run = p.add_run(text)
        run.bold = bold
        _apply_font(run)
        return p

    # ---- Dados ----
    def _fmt_cpf(s):
        s = _re.sub(r'\D', '', s)
        m = _re.match(r'(\d{3})(\d{3})(\d{3})(\d{2})', s)
        return f"{m.group(1)}.{m.group(2)}.{m.group(3)}-{m.group(4)}" if m else s

    def _fmt_cep(s):
        s = _re.sub(r'\D', '', s)
        m = _re.match(r'(\d{5})(\d{3})', s)
        return f"{m.group(1)}-{m.group(2)}" if m else s

    cidade  = d.get("cidade", "")
    estado  = d.get("estado", "PR")
    razao   = d.get("razao_social", "")
    cnpj    = d.get("cnpj", "")
    log     = d.get("logradouro", "")
    num     = d.get("numero", "s/nº")
    bairro  = d.get("bairro", "")
    cep_fmt = _fmt_cep(d.get("cep", ""))
    admin   = d.get("admin_nome", "")
    cpf_fmt = _fmt_cpf(d.get("admin_cpf", ""))
    cargo   = d.get("admin_cargo", "Sócio(a) Administrador(a)")
    data    = d.get("data_hoje", "")

    end_str = f"{log}, nº {num}, Bairro {bairro}, CEP {cep_fmt}"

    # ---- Título ----
    para("DECLARAÇÃO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=24)

    # ---- Corpo ----
    if tipo == "sem_atividade":
        corpo = (
            f"{razao}, inscrita no CNPJ sob o nº {cnpj}, com sede no município de "
            f"{cidade}/{estado}, {end_str}, neste ato representada pelo(a) {cargo} "
            f"{admin}, inscrito(a) no CPF nº {cpf_fmt}, DECLARA, para os devidos fins, "
            f"que no endereço empresarial não são exercidas atividades de qualquer natureza, "
            f"sendo somente ponto de referência."
        )
    else:
        corpo = (
            f"{razao}, pessoa jurídica de direito privado, inscrita no CNPJ/MF {cnpj}, "
            f"com sede no município de {cidade}/{estado}, {end_str}, neste ato representada "
            f"pelo(a) {cargo} {admin}, inscrito(a) no CPF {cpf_fmt}, DECLARA para os devidos "
            f"fins, sob responsabilidade e penas da lei, que não possui funcionários "
            f"registrados na empresa."
        )

    para(corpo, sa=12, indent=INDENT)

    if tipo == "sem_funcionarios":
        para("Por ser expressão de verdade e nada mais tendo a declarar, assina a presente.",
             sa=24, indent=INDENT)

    # ---- Local e data ----
    para(f"{cidade}/{estado}, {data}.", sb=24, sa=36)

    # ---- Assinatura ----
    para("_" * 50, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    para(razao, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    para(f"CNPJ: {cnpj}", align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    if admin:
        para(admin,             align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
        para(cargo,             align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
        para(f"CPF: {cpf_fmt}", align=WD_ALIGN_PARAGRAPH.CENTER, sa=0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@leads_api_bp.route("/<lead_id>/gerar-declaracao/<tipo>")
def gerar_declaracao(lead_id, tipo):
    """Gera e retorna a declaração como arquivo .docx para download."""
    if tipo not in ("sem_atividade", "sem_funcionarios"):
        abort(400)
    lead = db.get_lead(lead_id)
    if not lead:
        abort(404)

    ficha_data = None
    parent = db.get_lead_parent(lead_id)
    if lead.get("ficha_id"):
        from db import get_ficha as _get_ficha
        ficha_data = _get_ficha(int(lead["ficha_id"]))
    elif parent and parent.get("ficha_id"):
        from db import get_ficha as _get_ficha
        ficha_data = _get_ficha(int(parent["ficha_id"]))

    d = _build_decl_data(lead, ficha_data, parent)
    docx_bytes = _make_docx_declaration(tipo, d)

    nomes = {
        "sem_atividade":    "Declaracao_Sem_Atividade_No_Local.docx",
        "sem_funcionarios": "Declaracao_Sem_Funcionarios.docx",
    }
    return Response(
        docx_bytes,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{nomes[tipo]}"'},
    )


@leads_api_bp.route("/<lead_id>/gerar-documento-ia", methods=["POST"])
def gerar_documento_ia(lead_id):
    """Usa IA para gerar um documento personalizado com base nos dados do formulário."""
    import urllib.request as _ureq
    lead = db.get_lead(lead_id)
    if not lead:
        abort(404)

    data = request.get_json(silent=True) or {}
    contexto_usuario = (data.get("contexto") or "").strip()
    if not contexto_usuario:
        return jsonify({"error": "Informe o contexto do documento."}), 400

    api_key = os.environ.get("OPENAI_API_KEY", "")
    if not api_key:
        return jsonify({"error": "OPENAI_API_KEY não configurada."}), 500

    ficha_data = None
    parent = db.get_lead_parent(lead_id)
    if lead.get("ficha_id"):
        from db import get_ficha as _get_ficha
        ficha_data = _get_ficha(int(lead["ficha_id"]))
    elif parent and parent.get("ficha_id"):
        from db import get_ficha as _get_ficha
        ficha_data = _get_ficha(int(parent["ficha_id"]))

    d = _build_decl_data(lead, ficha_data, parent)

    import re as _re

    def _fmt_cpf(s):
        s = _re.sub(r'\D', '', s)
        m = _re.match(r'(\d{3})(\d{3})(\d{3})(\d{2})', s)
        return f"{m.group(1)}.{m.group(2)}.{m.group(3)}-{m.group(4)}" if m else s

    def _fmt_cep(s):
        s = _re.sub(r'\D', '', s)
        m = _re.match(r'(\d{5})(\d{3})', s)
        return f"{m.group(1)}-{m.group(2)}" if m else s

    end_fmt = (
        f"{d['logradouro']}, nº {d['numero']}"
        + (f", {d['bairro']}" if d['bairro'] else "")
        + (f", CEP {_fmt_cep(d['cep'])}" if d['cep'] else "")
    )

    dados_resumo = (
        f"Razão Social: {d['razao_social']}\n"
        f"CNPJ: {d['cnpj']}\n"
        f"Endereço: {end_fmt}\n"
        f"Município/UF: {d['cidade']}/{d['estado']}\n"
        f"Representante legal: {d['admin_nome']}, {d['admin_cargo']}, CPF {_fmt_cpf(d['admin_cpf'])}\n"
        f"Data de referência: {d['data_hoje']}"
    )

    system_prompt = """\
Você é um advogado empresarial sênior com 30 anos de experiência em Direito Empresarial brasileiro. \
Elabora declarações, requerimentos, ofícios e documentos societários com linguagem técnica impecável, \
observando rigorosamente as normas da ABNT e os padrões formais exigidos pelos órgãos públicos e registros comerciais.

REGRAS OBRIGATÓRIAS DE FORMATAÇÃO (siga-as à risca):

1. Estruture o documento usando exatamente estas marcações de seção em linhas separadas:
   == CABECALHO ==
   == TITULO ==
   == CORPO ==
   == FECHO ==
   == LOCAL_DATA ==
   == ASSINATURA ==

2. CABECALHO: dados da empresa centralizados (razão social em maiúsculas e negrito implícito, \
depois CNPJ, endereço, município/UF — cada item em linha separada).

3. TITULO: nome do documento em MAIÚSCULAS (ex: DECLARAÇÃO, REQUERIMENTO, OFÍCIO Nº ...).

4. CORPO: texto justificado com parágrafos separados por linha em branco. \
Primeira linha de cada parágrafo com recuo. \
Linguagem formal, técnica e objetiva. Cite artigos de lei quando pertinente.

5. FECHO:
   - Se for DECLARAÇÃO: feche com "Por ser verdade, firmo a presente declaração para que surta seus legais e jurídicos efeitos."
   - Se for REQUERIMENTO ou OFÍCIO: feche com "Nestes termos, pede deferimento." ou fecho equivalente ao tipo.
   - Se for outro tipo: use o fecho formal apropriado.

6. LOCAL_DATA: "{Cidade}/{UF}, {data por extenso}."

7. ASSINATURA: linha de traços (______...), razão social, CNPJ, nome do representante, cargo, CPF — \
cada item em linha separada, centralizado.

Retorne SOMENTE o documento estruturado com as marcações. Sem explicações, sem comentários adicionais."""

    user_prompt = (
        f"DADOS DA EMPRESA:\n{dados_resumo}\n\n"
        f"DOCUMENTO SOLICITADO:\n{contexto_usuario}\n\n"
        "Elabore o documento completo seguindo rigorosamente as regras de formatação."
    )

    payload = json.dumps({
        "model": "gpt-4o-mini",
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_prompt},
        ],
        "max_tokens": 1800,
        "temperature": 0.2,
    }).encode()

    try:
        req = _ureq.Request(
            "https://api.openai.com/v1/chat/completions",
            data=payload,
            headers={"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"},
        )
        with _ureq.urlopen(req, timeout=30) as resp:
            resultado = json.loads(resp.read())
            texto = resultado["choices"][0]["message"]["content"].strip()
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    # ------------------------------------------------------------------ #
    # Build .docx from structured AI output (section markers)            #
    # ------------------------------------------------------------------ #
    import io
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    FONT  = "Arial"
    SIZE  = Pt(12)
    SPACE = Pt(18)   # 1,5 × 12 pt
    IND   = Cm(1.25)

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(3)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(3)
        sec.right_margin  = Cm(2)

    def _apply_font(run, bold=False, size=None):
        run.font.name = FONT
        run.font.size = size or SIZE
        run.bold = bold
        rPr = run._r.get_or_add_rPr()
        for old in rPr.findall(qn('w:rFonts')):
            rPr.remove(old)
        rf = OxmlElement('w:rFonts')
        rf.set(qn('w:ascii'), FONT)
        rf.set(qn('w:hAnsi'), FONT)
        rf.set(qn('w:cs'),    FONT)
        rPr.insert(0, rf)

    def _para(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False,
               sb=0, sa=0, indent=None, size=None):
        p = doc.add_paragraph()
        p.alignment = align
        fmt = p.paragraph_format
        fmt.line_spacing  = SPACE
        fmt.space_before  = Pt(sb)
        fmt.space_after   = Pt(sa)
        if indent is not None:
            fmt.first_line_indent = indent
        run = p.add_run(text)
        _apply_font(run, bold=bold, size=size)
        return p

    # Parse sections from AI output
    SECTION_RE = _re.compile(r'^==\s*(CABECALHO|TITULO|CORPO|FECHO|LOCAL_DATA|ASSINATURA)\s*==\s*$',
                             _re.IGNORECASE)
    sections: dict[str, list[str]] = {}
    current = "CORPO"
    for line in texto.split("\n"):
        m = SECTION_RE.match(line.strip())
        if m:
            current = m.group(1).upper()
            sections.setdefault(current, [])
        else:
            sections.setdefault(current, []).append(line)

    def _lines(key):
        return [l.rstrip() for l in sections.get(key, [])]

    def _non_empty(key):
        return [l for l in _lines(key) if l.strip()]

    # ---- CABEÇALHO (centralizado) ----
    cab_lines = _non_empty("CABECALHO")
    if cab_lines:
        # First line = razão social (bold, slightly larger)
        _para(cab_lines[0], align=WD_ALIGN_PARAGRAPH.CENTER,
              bold=True, size=Pt(12), sb=0, sa=2)
        for cl in cab_lines[1:]:
            _para(cl, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
        # Separator line
        _para("_" * 60, align=WD_ALIGN_PARAGRAPH.CENTER, sa=12)
    else:
        # Fallback: print company data from d dict
        _para(d.get("razao_social","").upper(),
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, sa=2)
        _para(f"CNPJ: {d.get('cnpj','')}",
              align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
        _para(f"{d.get('logradouro','')}, nº {d.get('numero','')}, {d.get('bairro','')}",
              align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
        _para(f"{d.get('cidade','')}/{d.get('estado','')}",
              align=WD_ALIGN_PARAGRAPH.CENTER, sa=12)
        _para("_" * 60, align=WD_ALIGN_PARAGRAPH.CENTER, sa=12)

    # ---- TÍTULO ----
    titulo_lines = _non_empty("TITULO")
    titulo_text = " ".join(titulo_lines).strip() if titulo_lines else "DOCUMENTO"
    _para(titulo_text.upper(), align=WD_ALIGN_PARAGRAPH.CENTER,
          bold=True, sb=6, sa=24)

    # ---- CORPO ----
    corpo_lines = _lines("CORPO")
    buf_para: list[str] = []
    def _flush_para():
        text = " ".join(buf_para).strip()
        if text:
            _para(text, indent=IND, sb=0, sa=6)
        buf_para.clear()

    for line in corpo_lines:
        if line.strip() == "":
            _flush_para()
        else:
            buf_para.append(line.strip())
    _flush_para()

    # ---- FECHO ----
    fecho_lines = _lines("FECHO")
    buf_para2: list[str] = []
    def _flush_fecho():
        text = " ".join(buf_para2).strip()
        if text:
            _para(text, indent=IND, sb=6, sa=6)
        buf_para2.clear()

    for line in fecho_lines:
        if line.strip() == "":
            _flush_fecho()
        else:
            buf_para2.append(line.strip())
    _flush_fecho()

    # ---- LOCAL E DATA ----
    local_data_lines = _non_empty("LOCAL_DATA")
    local_text = " ".join(local_data_lines).strip()
    if local_text:
        _para(local_text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=18, sa=24)

    # ---- ASSINATURA ----
    assin_lines = _non_empty("ASSINATURA")
    if assin_lines:
        for al in assin_lines:
            stripped = al.strip()
            is_line = set(stripped) <= set("_- ")
            _para(stripped if stripped else " ",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  bold=(not is_line and stripped == stripped.upper() and len(stripped) > 4),
                  sa=2)
    else:
        # Fallback signature block
        _para("_" * 50, align=WD_ALIGN_PARAGRAPH.CENTER, sb=6, sa=2)
        _para(d.get("razao_social",""), align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, sa=2)
        _para(f"CNPJ: {d.get('cnpj','')}", align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
        if d.get("admin_nome"):
            _para(d["admin_nome"], align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
            _para(d.get("admin_cargo",""), align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
            _para(f"CPF: {_fmt_cpf(d.get('admin_cpf',''))}", align=WD_ALIGN_PARAGRAPH.CENTER, sa=0)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    # Save to storage + create file record + add comment
    docx_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    import re as _re2
    safe_name = _re2.sub(r'[^a-z0-9_]', '_', contexto_usuario[:40].lower()).strip('_') or 'documento'
    filename = f"IA_{safe_name}.docx"
    try:
        storage = get_storage()
        storage_key, size = storage.save(lead_id, filename, buf, docx_mime)
        file_id = db.add_file(
            lead_id,
            filename=filename,
            storage_key=storage_key,
            size_bytes=size,
            mime_type=docx_mime,
        )
        actor = session.get("user_name") or "Sistema"
        db.add_comment(
            lead_id,
            body=f"Documento gerado por IA: {contexto_usuario}",
            author=actor,
            attachment_key=storage_key,
            attachment_name=filename,
            attachment_mime=docx_mime,
        )
    except Exception:
        file_id = None

    return jsonify({
        "ok": True,
        "file_id": file_id,
        "filename": filename,
    })


# ---------------------------------------------------------------------------
# Proxies IBGE (cidade + CNAE) — evita CORS e centraliza cache
# ---------------------------------------------------------------------------

_city_cache: dict[str, list[str]] = {}
_cnae_cache: list[dict] | None = None


def _http_get_json(url: str, timeout: int = 10):
    req = urllib.request.Request(url, headers={"User-Agent": "contratos-legal/1.0"})
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        return json.loads(resp.read().decode("utf-8"))


@leads_api_bp.route("/ibge/cidades/<uf>")
def ibge_cidades(uf):
    uf = uf.upper().strip()
    if uf in _city_cache:
        return jsonify(_city_cache[uf])
    try:
        url = (f"https://servicodados.ibge.gov.br/api/v2/localidades/estados/"
               f"{quote(uf)}/municipios?orderBy=nome")
        data = _http_get_json(url)
        names = [m["nome"] for m in data]
        _city_cache[uf] = names
        return jsonify(names)
    except Exception as e:
        return jsonify({"error": str(e)}), 502


# ---------------------------------------------------------------------------
# Notifications
# ---------------------------------------------------------------------------

def _notify_by_name(responsible_name: str, lead_id: str, notif_type: str,
                    message: str, actor_name: str | None = None) -> None:
    """Creates a notification for the first user matching responsible_name."""
    users = db.list_users()
    for u in users:
        if u["name"] == responsible_name:
            db.create_notification(u["id"], lead_id, notif_type, message, actor_name)
            break


@leads_api_bp.route("/users-list")
def users_list():
    """Returns list of users for @mention autocomplete."""
    users = db.list_users()
    return jsonify([{"id": u["id"], "name": u["name"]} for u in users])


@leads_api_bp.route("/notifications")
def get_notifications():
    user_id = session.get("user_id")
    if not user_id:
        return jsonify([])
    notifications = db.list_notifications(user_id)
    unread = db.count_unread_notifications(user_id)
    return jsonify({"notifications": notifications, "unread": unread})


@leads_api_bp.route("/notifications/read-all", methods=["POST"])
def notifications_read_all():
    user_id = session.get("user_id")
    if user_id:
        db.mark_all_notifications_read(user_id)
    return jsonify({"ok": True})


@leads_api_bp.route("/notifications/<notif_id>/read", methods=["POST"])
def notification_read(notif_id):
    db.mark_notification_read(notif_id)
    return jsonify({"ok": True})


@leads_api_bp.route("/ibge/cnaes")
def ibge_cnaes():
    global _cnae_cache
    if _cnae_cache is not None:
        return jsonify(_cnae_cache)
    try:
        data = _http_get_json("https://servicodados.ibge.gov.br/api/v2/cnae/subclasses")
        items = []
        for it in data:
            cid = str(it["id"])
            # 0111301 -> 0111-3/01
            formatted = f"{cid[:4]}-{cid[4]}/{cid[5:]}" if len(cid) == 7 else cid
            items.append({
                "id": cid,
                "code": formatted,
                "label": f"{formatted} — {it['descricao']}",
                "descricao": it["descricao"],
            })
        _cnae_cache = items
        return jsonify(items)
    except Exception as e:
        return jsonify({"error": str(e)}), 502
