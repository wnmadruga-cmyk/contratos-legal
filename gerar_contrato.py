#!/usr/bin/env python3
"""
Gerador de Contrato Social LTDA
Uso: python3 gerar_contrato.py <ficha.json> [saida.docx]
"""

import json
import sys
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from num2words import num2words

# ---------------------------------------------------------------------------
# Constantes de formatação (retiradas do template original)
# ---------------------------------------------------------------------------

FONT_NAME   = "Arial"
FONT_SIZE   = Pt(12)
FONT_SIZE_N = 152400  # EMU equivalente a 12pt

# Indentações em EMU (convertidas do template)
INDENT_NORMAL = -3175    # padrão do documento
INDENT_SOCIO  = 450215   # recuo das qualificações dos sócios

# Espaçamentos após parágrafo
SPC_NORMAL  = 76200   # ~6pt
SPC_CLAUSULA = 0
SPC_CNAE    = 22225
SPC_CAPITAL = 71755
SPC_ABERTURA = 66675
SPC_DATA    = 311785
SPC_ASSINATURA = 53975

# ---------------------------------------------------------------------------
# Mapeamentos
# ---------------------------------------------------------------------------

MESES_PT = {
    1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril",
    5: "maio", 6: "junho", 7: "julho", 8: "agosto",
    9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
}

REGIME_BENS = {
    "comunhao_parcial":   "comunhão parcial de bens",
    "comunhao_universal": "comunhão universal de bens",
    "separacao_total":    "separação total de bens",
    "participacao_final": "participação final nos aquestos",
}

FORMA_INTEGRALIZACAO = {
    "dinheiro":        "em moeda corrente no País",
    "moeda":           "em moeda corrente no País",
    "moeda corrente":  "em moeda corrente no País",
    "moeda_corrente":  "em moeda corrente no País",
    "movel":           "em bens móveis",
    "bem movel":       "em bens móveis",
    "bens moveis":     "em bens móveis",
    "bens_moveis":     "em bens móveis",
    "imovel":          "em bens imóveis",
    "bem imovel":      "em bens imóveis",
    "bens imoveis":    "em bens imóveis",
    "bens_imoveis":    "em bens imóveis",
}

ESTADO_CIVIL = {
    "solteiro":   ("solteiro",   "solteira"),
    "casado":     ("casado",     "casada"),
    "divorciado": ("divorciado", "divorciada"),
    "viuvo":      ("viúvo",      "viúva"),
    "separado":   ("separado",   "separada"),
}

# ---------------------------------------------------------------------------
# Utilitários de texto
# ---------------------------------------------------------------------------

def formatar_cpf(cpf: str) -> str:
    cpf = "".join(filter(str.isdigit, cpf))
    if len(cpf) == 11:
        return f"{cpf[:3]}.{cpf[3:6]}.{cpf[6:9]}-{cpf[9:]}"
    return cpf


def formatar_cep(cep: str) -> str:
    cep = "".join(filter(str.isdigit, cep))
    if len(cep) == 8:
        return f"{cep[:5]}-{cep[5:]}"
    return cep


def valor_por_extenso(valor: float) -> str:
    inteiro = int(valor)
    extenso = num2words(inteiro, lang="pt_BR")
    return f"{extenso} {'real' if inteiro == 1 else 'reais'}"


def cotas_por_extenso(qtd: int) -> str:
    return num2words(int(qtd), lang="pt_BR")


def formatar_valor_reais(valor: float) -> str:
    return f"{valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def formatar_data_por_extenso(dt: datetime) -> str:
    return f"{dt.day:02d} de {MESES_PT[dt.month]} de {dt.year}"


def inferir_genero(socio: dict) -> str:
    """Retorna 'm' ou 'f'. Usa campo 'genero' do JSON se existir, senão infere pela profissão."""
    g = socio.get("genero", "")
    if g:
        return "f" if g.strip().lower()[0] == "f" else "m"
    prof = socio.get("profissao", "").strip().lower()
    if prof.endswith("a") and not prof.endswith("ista"):
        return "f"
    return "m"


def normalizar_profissao(prof: str) -> str:
    """Profissão sempre em minúsculo, independente de como vier no formulário."""
    return prof.strip().lower()


def get_integralizacoes_socio(socio: dict) -> list:
    """
    Retorna lista de integralizações do sócio.
    Suporta novo formato (integralizacoes[]) e formato antigo (observacoesIntegralizacao).
    """
    if "integralizacoes" in socio and socio["integralizacoes"]:
        return socio["integralizacoes"]
    obs = socio.get("observacoesIntegralizacao", "").strip()
    val = int(socio.get("quantidadeCotas", 0)) * float(socio.get("valorUnitarioCota", 1))
    if obs:
        return [{"tipo": obs, "valor": val, "descricao": ""}]
    return [{"tipo": "moeda", "valor": val, "descricao": ""}]


def texto_tipo_integralizacao(tipo: str) -> str:
    return FORMA_INTEGRALIZACAO.get(tipo.strip().lower(), tipo)


def resolver_forma_integralizacao(socios: list, empresa: dict) -> str:
    """
    Retorna texto resumido para o parágrafo do capital.
    Se todos os sócios integralizam da mesma forma → usa essa forma.
    Se misto → texto genérico (detalhes vêm no parágrafo após a tabela).
    """
    forma_empresa = empresa.get("formaIntegralizacao", "").strip().lower()
    if forma_empresa:
        return texto_tipo_integralizacao(forma_empresa)

    formas = set()
    for s in socios:
        for item in get_integralizacoes_socio(s):
            formas.add(texto_tipo_integralizacao(item.get("tipo", "moeda")))

    if len(formas) == 1:
        return formas.pop()
    return "conforme especificado abaixo"


def estado_civil_texto(socio: dict, genero: str) -> str:
    ec_key  = socio.get("estadoCivil", "solteiro").lower()
    idx     = 0 if genero == "m" else 1
    ec_pair = ESTADO_CIVIL.get(ec_key, (ec_key, ec_key))
    ec_str  = ec_pair[idx]

    regime_key = socio.get("regimeBens", "")
    if ec_key == "casado" and regime_key:
        regime = REGIME_BENS.get(regime_key, regime_key.replace("_", " "))
        ec_str = f"{ec_str}, sob o regime de {regime}"

    # União estável: append if marked
    uniao = socio.get("uniaoEstavel") or socio.get("uniao_estavel")
    companheiro = (socio.get("nomeCompanheiro") or socio.get("nome_companheiro") or "").strip()
    if uniao and str(uniao).lower() not in ("false", "0", ""):
        if companheiro:
            ec_str = f"{ec_str}, convivente em união estável com {companheiro}"
        else:
            ec_str = f"{ec_str}, convivente em união estável"

    return ec_str


def title_case(texto: str) -> str:
    """Title Case respeitando palavras já acentuadas; normaliza maiúsculas/minúsculas."""
    return texto.strip().title()


def formatar_endereco(end: dict) -> str:
    tipo      = title_case(end.get("logradouroTipo", ""))
    desc      = title_case(end.get("logradouroDescricao", ""))
    num       = end.get("numero", "").strip()
    compl     = end.get("complemento", "").strip()
    lote      = end.get("lote", "").strip()
    quadra    = end.get("quadra", "").strip()
    inscricao = end.get("inscricaoImobiliaria", "").strip()
    bairro    = title_case(end.get("bairro", ""))
    cidade    = title_case(end.get("cidade", ""))
    estado    = end.get("estado", "").strip().upper()
    cep       = formatar_cep(end.get("cep", ""))

    # Evita duplicar o tipo (ex: "Rua Rua João Pessoa")
    if desc.lower().startswith(tipo.lower() + " "):
        logradouro = desc
    else:
        logradouro = f"{tipo} {desc}".strip()

    estado_nome = ESTADOS_NOMES.get(estado, estado)

    partes = [f"{logradouro}, nº {num}"]
    if compl:
        partes.append(compl)
    partes.append(bairro)
    partes.append(f"cidade de {cidade} – {estado_nome}")
    partes.append(f"CEP: {cep}")
    return ", ".join(partes)


def nacionalidade(genero: str, socio: dict) -> str:
    """Retorna a nacionalidade concordando com o gênero.
    Aceita qualquer forma (masculino, feminino, sem acento) e
    retorna a forma correta para o gênero informado.
    """
    import unicodedata
    nac = socio.get("nacionalidade", "").strip().lower()
    if not nac:
        nac = "brasileiro"

    idx = 0 if genero == "m" else 1

    # Tenta direto
    par = NACIONALIDADES_MAP.get(nac)
    if par:
        return par[idx]

    # Tenta sem acento
    nac_ascii = unicodedata.normalize("NFD", nac).encode("ascii", "ignore").decode().lower()
    par = NACIONALIDADES_MAP.get(nac_ascii)
    if par:
        return par[idx]

    # Fallback: retorna o valor como digitado
    return nac


ESTADOS_NOMES = {
    "AC": "Acre",              "AL": "Alagoas",         "AP": "Amapá",
    "AM": "Amazonas",          "BA": "Bahia",           "CE": "Ceará",
    "DF": "Distrito Federal",  "ES": "Espírito Santo",  "GO": "Goiás",
    "MA": "Maranhão",          "MT": "Mato Grosso",     "MS": "Mato Grosso do Sul",
    "MG": "Minas Gerais",      "PA": "Pará",            "PB": "Paraíba",
    "PR": "Paraná",            "PE": "Pernambuco",      "PI": "Piauí",
    "RJ": "Rio de Janeiro",    "RN": "Rio Grande do Norte", "RS": "Rio Grande do Sul",
    "RO": "Rondônia",          "RR": "Roraima",         "SC": "Santa Catarina",
    "SP": "São Paulo",         "SE": "Sergipe",         "TO": "Tocantins",
}

# Cada entrada mapeia AMBAS as formas (masc. e fem.) para o par (masc., fem.)
_NAC_PARES = [
    ("brasileiro",   "brasileira"),
    ("americano",    "americana"),
    ("argentino",    "argentina"),
    ("italiano",     "italiana"),
    ("espanhol",     "espanhola"),
    ("português",    "portuguesa"),
    ("alemão",       "alemã"),
    ("francês",      "francesa"),
    ("japonês",      "japonesa"),
    ("chinês",       "chinesa"),
    ("uruguaio",     "uruguaia"),
    ("paraguaio",    "paraguaia"),
    ("colombiano",   "colombiana"),
    ("peruano",      "peruana"),
    ("boliviano",    "boliviana"),
    ("venezuelano",  "venezuelana"),
    ("chileno",      "chilena"),
    ("mexicano",     "mexicana"),
    ("cubano",       "cubana"),
    ("equatoriano",  "equatoriana"),
    ("panamenho",    "panamenha"),
    ("hondurenho",   "hondurenha"),
    ("guatemalteco", "guatemalteca"),
    ("salvadorenho", "salvadorenha"),
    ("costarricense","costarricense"),
    ("dominicano",   "dominicana"),
    ("haitiano",     "haitiana"),
    ("nigeriano",    "nigeriana"),
    ("angolano",     "angolana"),
    ("moçambicano",  "moçambicana"),
    ("cabo-verdiano","cabo-verdiana"),
]

NACIONALIDADES_MAP: dict = {}
for _m, _f in _NAC_PARES:
    NACIONALIDADES_MAP[_m] = (_m, _f)   # chave masculina
    NACIONALIDADES_MAP[_f] = (_m, _f)   # chave feminina
    # versões sem acento para robustez
    def _sem_acento(s):
        import unicodedata
        return unicodedata.normalize("NFD", s).encode("ascii","ignore").decode()
    NACIONALIDADES_MAP[_sem_acento(_m)] = (_m, _f)
    NACIONALIDADES_MAP[_sem_acento(_f)] = (_m, _f)

DOCUMENTO_TIPO = {
    "rg":                   ("RG",         "o",  "expedido"),
    "cnh":                  ("CNH",        "a",  "expedida"),
    "passaporte":           ("Passaporte", "o",  "expedido"),
    "carteira_profissional": ("Carteira Profissional", "a", "expedida"),
    "rne":                  ("RNE",        "o",  "expedido"),
    "ctps":                 ("CTPS",       "a",  "expedida"),
}


def formatar_documento(socio: dict, genero: str) -> str:
    """Retorna texto do documento de identificação ou string vazia se não houver."""
    doc = socio.get("documentoIdentificacao", {})
    if not doc or not doc.get("numero"):
        return ""

    tipo_key   = doc.get("tipo", "").strip().lower()
    numero     = doc.get("numero", "").strip()
    orgao      = doc.get("orgaoExpedidor", "").strip().upper()
    data_exp   = doc.get("dataExpedicao", "").strip()

    nome_doc, artigo, particip = DOCUMENTO_TIPO.get(tipo_key, (tipo_key.upper(), "o", "expedido"))

    portador_a = "portador" if genero == "m" else "portadora"

    partes = [f"{portador_a} d{artigo} {nome_doc} nº {numero}"]
    if orgao:
        partes.append(f"{particip} pel{artigo} {orgao}")
    if data_exp:
        try:
            dt = datetime.strptime(data_exp, "%Y-%m-%d")
            partes.append(f"em {dt.strftime('%d/%m/%Y')}")
        except ValueError:
            partes.append(f"em {data_exp}")

    return ", ".join(partes)


def socio_qualificado(socio: dict) -> str:
    genero    = inferir_genero(socio)
    cpf       = formatar_cpf(socio.get("cpf", ""))
    nac       = nacionalidade(genero, socio)
    ec        = estado_civil_texto(socio, genero)
    prof      = normalizar_profissao(socio.get("profissao", ""))
    nascido_a = "nascido" if genero == "m" else "nascida"
    resid_a   = "residente e domiciliad" + ("o" if genero == "m" else "a")

    nasc = socio.get("dataNascimento", "")
    if nasc:
        try:
            dt = datetime.strptime(nasc, "%Y-%m-%d")
            nasc_txt = f", {nascido_a} em {dt.strftime('%d/%m/%Y')}"
        except ValueError:
            nasc_txt = f", {nascido_a} em {nasc}"
    else:
        nasc_txt = f", {nascido_a} em ___/___/______"

    doc_txt = formatar_documento(socio, genero)
    doc_parte = f", {doc_txt}" if doc_txt else ""

    end = formatar_endereco(socio.get("endereco", {}))
    qualif = f"{nac}, {ec}, {prof}{nasc_txt}, n° do CPF {cpf}{doc_parte}, {resid_a} na {end}"

    # Representante legal — quando menor ou legalmente incapaz
    if socio.get("menorOuIncapaz") and socio.get("representante"):
        rep = socio["representante"]
        rep_gen = inferir_genero(rep)
        rep_tipo_raw = rep.get("tipo", "pai").lower()
        if rep_tipo_raw == "mae":
            rep_label = "sua mãe"
        elif rep_tipo_raw == "tutor":
            rep_label = "seu tutor" if rep_gen == "m" else "sua tutora"
        elif rep_tipo_raw == "curador":
            rep_label = "seu curador" if rep_gen == "m" else "sua curadora"
        else:
            rep_label = "seu pai"
        rep_qualif = socio_qualificado(rep)  # recursivo — rep não terá menorOuIncapaz
        repr_o_a   = "o" if genero == "m" else "a"
        qualif += (f", representad{repr_o_a} neste ato por {rep_label}, "
                   f"{rep.get('nome', '').upper()}, {rep_qualif}")

    return qualif


# ---------------------------------------------------------------------------
# Helpers de formatação de parágrafo
# ---------------------------------------------------------------------------

def set_paragraph_format(para,
                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          left_indent=INDENT_NORMAL,
                          first_line_indent=None,
                          space_before=None,
                          space_after=SPC_NORMAL):
    pf = para.paragraph_format
    pf.alignment       = alignment
    pf.left_indent     = left_indent
    pf.first_line_indent = first_line_indent
    if space_before is not None:
        pf.space_before = space_before
    pf.space_after     = space_after


def add_run(para, text: str, bold=False):
    run = para.add_run(text)
    run.bold      = bold
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    return run


def set_table_width(table, width_dxa: int):
    """Define a largura total da tabela em dxa."""
    tblPr = table._tbl.find(qn("w:tblPr"))
    tblW  = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(width_dxa))
    tblW.set(qn("w:type"), "dxa")
    existing = tblPr.find(qn("w:tblW"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblW)


def set_col_width(cell, width_dxa: int):
    tc   = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(width_dxa))
    tcW.set(qn("w:type"), "dxa")
    existing = tcPr.find(qn("w:tcW"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcW)


def format_cell_text(cell, text: str, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    add_run(p, text, bold=bold)


def remover_bordas_tabela(table):
    """Remove todas as bordas visíveis de uma tabela."""
    tblPr = table._tbl.find(qn("w:tblPr"))
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


# ---------------------------------------------------------------------------
# Cabeçalho do documento
# ---------------------------------------------------------------------------

def _add_page_number(para):
    """Adiciona campo 'Página X de Y' a um parágrafo."""
    def _field(p, instrucao):
        run = p.add_run()
        run.font.name = FONT_NAME
        run.font.size = Pt(9)
        fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin')
        it  = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = f' {instrucao} '
        fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'end')
        run.element.append(fc1); run.element.append(it); run.element.append(fc2)

    r = para.add_run("Página "); r.font.name = FONT_NAME; r.font.size = Pt(9)
    _field(para, "PAGE")
    r = para.add_run(" de ");   r.font.name = FONT_NAME; r.font.size = Pt(9)
    _field(para, "NUMPAGES")


def adicionar_cabecalho(doc, razao_social: str):
    section = doc.sections[0]
    header  = section.header
    for p in header.paragraphs:
        p.clear()

    # Linha 1: título centralizado em negrito
    p1 = header.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p1, "CONTRATO SOCIAL DE CONSTITUIÇÃO DE SOCIEDADE EMPRESÁRIA LIMITADA", bold=True)

    # Linha 2: razão social centralizada
    p2 = header.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, razao_social)

    # Linha 3: numeração de páginas (direita)
    p3 = header.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.paragraph_format.space_before = None
    p3.paragraph_format.space_after  = None
    _add_page_number(p3)


# ---------------------------------------------------------------------------
# Construção do documento
# ---------------------------------------------------------------------------

def gerar_contrato(dados: dict, caminho_saida):
    empresa = dados["empresa"]
    socios  = empresa["socios"]
    resumo  = dados.get("resumo", {})

    ts = dados.get("timestamp", "")
    try:
        data_contrato = datetime.fromisoformat(ts.replace("Z", "+00:00"))
    except Exception:
        data_contrato = datetime.today()

    data_fmt     = formatar_data_por_extenso(data_contrato)
    capital      = empresa.get("capitalSocial", 0)
    total_cotas  = sum(s.get("quantidadeCotas", 0) for s in socios)
    razao_social = empresa.get("razaoSocial", "").upper()
    objeto_social = empresa.get("objetoSocial", "")
    # Filtra atividades vazias (sem CNAE e sem descrição)
    atividades   = [a for a in empresa.get("atividades", [])
                    if a.get("cnae", "").strip() or a.get("descricao", "").strip()]
    end_comercial = empresa.get("enderecoComercial", {})
    administradores = [s for s in socios if s.get("administrador")]

    # Modo de administração: empresa > primeiro admin > padrão isolada
    tipo_adm = empresa.get("tipoAdministracao", "")
    if not tipo_adm and administradores:
        tipo_adm = administradores[0].get("tipoAdministracao", "isolada")
    tipo_adm = tipo_adm.lower() if tipo_adm else "isolada"

    # Microempresa: campo booleano ou classificação (me / microempresa / epp)
    classificacao = empresa.get("classificacao", "").lower()
    microempresa  = empresa.get("microempresa",
                                classificacao in ("me", "epp") or
                                classificacao.startswith("micro") or
                                classificacao.startswith("epp"))

    # Unipessoal: 1 sócio
    is_unipessoal = len(socios) == 1

    # Helper para buscar cláusula do DB
    from db import get_clausula_texto as _get_cl

    # -----------------------------------------------------------------------
    doc = Document()

    # Margens conforme template original
    for section in doc.sections:
        section.top_margin    = Cm(2.88)
        section.bottom_margin = Cm(2.50)
        section.left_margin   = Cm(1.99)
        section.right_margin  = Cm(1.75)

    # Remover parágrafo inicial vazio que o python-docx cria
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # Cabeçalho
    adicionar_cabecalho(doc, razao_social)

    # -----------------------------------------------------------------------
    # Parágrafo de abertura
    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=SPC_ABERTURA, left_indent=INDENT_NORMAL)
    add_run(p, "Pelo presente instrumento particular de Contrato Social:")

    # Qualificação dos sócios (recuado) — nome sempre em caixa alta
    for socio in socios:
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_SOCIO, space_after=SPC_NORMAL)
        add_run(p, socio["nome"].upper(), bold=True)
        add_run(p, ", " + socio_qualificado(socio) + ";")

    # Resolvem / Resolve (mesmo recuo dos sócios)
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_SOCIO, space_after=SPC_NORMAL)
    if is_unipessoal:
        add_run(p, (
            "Resolve constituir uma sociedade empresária limitada, "
            "nos termos da Lei n° 10.406/2002, mediante as condições e cláusulas seguintes:"
        ))
    else:
        add_run(p, (
            "Resolvem, em comum acordo, constituir uma sociedade empresária limitada, "
            "nos termos da Lei n° 10.406/2002, mediante as condições e cláusulas seguintes:"
        ))

    # --- CLÁUSULA I ---
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
    add_run(p, "CLÁUSULA I - DO NOME EMPRESARIAL (art. 997, II, CC)", bold=True)

    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
    add_run(p, "A sociedade adotará como nome empresarial: ")
    add_run(p, razao_social + ".", bold=True)

    # --- CLÁUSULA II ---
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
    add_run(p, "CLÁUSULA II - DA SEDE (art. 997, II, CC)", bold=True)

    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=0, first_line_indent=0, space_after=SPC_NORMAL)
    add_run(p, "A sociedade terá sua sede no seguinte endereço: ")
    add_run(p, formatar_endereco(end_comercial) + ";", bold=True)

    # --- CLÁUSULA III ---
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
    add_run(p, "CLÁUSULA III - DO OBJETO SOCIAL (art. 997, II, CC)", bold=True)

    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
    add_run(p, "A sociedade terá por objeto o exercício das seguintes atividades econômicas: ")
    add_run(p, objeto_social.upper(), bold=True)

    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
    add_run(p, "E exercerá as seguintes atividades:")

    for atv in atividades:
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CNAE)
        desc_raw = atv.get('descricao', '').strip()
        desc_fmt = desc_raw.capitalize() if desc_raw else desc_raw
        add_run(p, f"CNAE Nº {atv['cnae']} - {desc_fmt}")

    # --- CLÁUSULA IV ---
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
    add_run(p, "CLÁUSULA IV - DO INÍCIO DAS ATIVIDADES E PRAZO DE DURAÇÃO "
               "(art. 53, III, F, Decreto n° 1.800/96) ", bold=True)
    add_run(p, "A sociedade iniciará suas atividades em ")
    add_run(p, data_contrato.strftime("%d/%m/%Y"), bold=True)
    add_run(p, " e seu prazo de duração será por tempo indeterminado.")

    # --- CLÁUSULA V ---
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
    add_run(p, "CLÁUSULA V - DO CAPITAL (ART. 997, III e IV e ART. 1.052 e 1.055, CC)", bold=True)

    capital_fmt  = formatar_valor_reais(capital)
    capital_ext  = valor_por_extenso(capital)
    cotas_int    = int(total_cotas)
    cotas_fmt    = f"{cotas_int:,}".replace(",", ".")
    cotas_ext    = cotas_por_extenso(cotas_int)

    # Parágrafo principal — sem mencionar a forma aqui
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
    add_run(p, (
        f"O capital será de R$ {capital_fmt} ({capital_ext}), "
        f"dividido em {cotas_fmt} ({cotas_ext}) quotas, "
        f"no valor nominal de R$ 1,00 (um real) cada uma, integralizado da seguinte forma:"
    ))

    # Detalhes de integralização por sócio (logo após o caput)
    for socio in socios:
        items = get_integralizacoes_socio(socio)
        for item in items:
            tipo_txt  = texto_tipo_integralizacao(item.get("tipo", "moeda"))
            val_item  = float(item.get("valor", 0))
            descricao = item.get("descricao", "").strip()
            nome_up   = socio["nome"].upper()
            val_fmt   = formatar_valor_reais(val_item)
            val_ext   = valor_por_extenso(val_item)

            linha = f"{nome_up} integraliza R$ {val_fmt} ({val_ext}) {tipo_txt}"
            if descricao:
                linha += f", sendo: {descricao}"
            linha += "."

            p = doc.add_paragraph()
            set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CNAE)
            add_run(p, linha)

    # Parágrafo único com a tabela de distribuição
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
    add_run(p, "Parágrafo único.", bold=True)
    if is_unipessoal:
        add_run(p, " O capital encontra-se subscrito e integralizado pelo sócio da seguinte forma:")
    else:
        add_run(p, " O capital encontra-se subscrito e integralizado pelos sócios da seguinte forma:")

    # Tabela de cotas — larguras exatas do template original
    # Colunas (dxa): Sócios=5118, %=1138, Quotas=1274, Valor=1844  Total=9374
    COL_WIDTHS = [5118, 1138, 1274, 1844]
    TABLE_WIDTH = 9374

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_width(table, TABLE_WIDTH)

    headers = ["Sócios", "%", "Quotas", "Valor em R$"]
    for i, (cell, txt) in enumerate(zip(table.rows[0].cells, headers)):
        set_col_width(cell, COL_WIDTHS[i])
        format_cell_text(cell, txt, bold=True)

    percentuais = {
        "".join(filter(str.isdigit, p["cpf"])): p
        for p in resumo.get("percentualPorSocio", [])
    }

    total_pct   = 0.0
    total_valor = 0.0

    for socio in socios:
        cpf_key  = "".join(filter(str.isdigit, socio["cpf"]))
        pct_data = percentuais.get(cpf_key, {})
        pct      = float(pct_data.get("percentual", 0))
        qtd      = int(socio.get("quantidadeCotas", 0))
        val      = qtd * float(socio.get("valorUnitarioCota", 1))
        total_pct   += pct
        total_valor += val

        row = table.add_row().cells
        for i, (cell, txt) in enumerate(zip(row, [
            socio["nome"].upper(),
            f"{pct:.2f}".replace(".", ","),
            f"{qtd:,}".replace(",", "."),
            formatar_valor_reais(val)
        ])):
            set_col_width(cell, COL_WIDTHS[i])
            format_cell_text(cell, txt)

    # Linha TOTAL
    row = table.add_row().cells
    for i, (cell, txt) in enumerate(zip(row, [
        "TOTAL:",
        f"{total_pct:.2f}".replace(".", ","),
        f"{cotas_int:,}".replace(",", "."),
        formatar_valor_reais(total_valor)
    ])):
        set_col_width(cell, COL_WIDTHS[i])
        format_cell_text(cell, txt, bold=True)

    # --- CLÁUSULA VI ---
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL,
                         space_before=SPC_NORMAL, space_after=SPC_CLAUSULA)
    add_run(p, "CLÁUSULA VI - DA ADMINISTRAÇÃO (ART. 997, VI; 1.013, 1.015; 1.064, CC)", bold=True)

    nomes_admin = [s["nome"] for s in administradores]

    if len(administradores) == 1:
        adm    = administradores[0]
        gen    = inferir_genero(adm)
        artigo = "pela sócia" if gen == "f" else "pelo sócio"
        modo   = "ISOLADAMENTE"

        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CAPITAL)
        add_run(p, f"A administração da sociedade será exercida {artigo} ")
        add_run(p, adm["nome"].upper(), bold=True)
        add_run(p, " que representará legalmente a sociedade ")
        add_run(p, modo, bold=True)
        add_run(p, " e poderá praticar todo e qualquer ato de gestão pertinente ao objeto social.")

    else:
        if tipo_adm == "conjunta":
            modo      = "CONJUNTAMENTE"
            sep_nomes = " e "
        elif tipo_adm == "isolada_conjunta":
            modo      = "ISOLADAMENTE e/ou CONJUNTAMENTE"
            sep_nomes = " e/ou "
        else:  # isolada
            modo      = "ISOLADAMENTE"
            sep_nomes = " e/ou "

        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CAPITAL)
        add_run(p, "A administração da sociedade será exercida pelos sócios ")
        for i, adm in enumerate(administradores):
            add_run(p, adm["nome"].upper(), bold=True)
            if i < len(administradores) - 2:
                add_run(p, ", ")
            elif i == len(administradores) - 2:
                add_run(p, sep_nomes)
        add_run(p, ", que representarão legalmente a sociedade ")
        add_run(p, modo, bold=True)
        add_run(p, " e poderão praticar todo e qualquer ato de gestão pertinente ao objeto social.")

    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
    add_run(p, "Parágrafo único.", bold=True)
    add_run(p, " Não constituindo o objeto social, a alienação ou a oneração de bens imóveis "
               "depende de autorização da maioria.")

    # --- CLÁUSULAS VII, VIII, IX (via DB) ---
    for cod in ("vii", "viii", "ix"):
        titulo_cl, corpo_cl = _get_cl(cod, is_unipessoal)
        if not titulo_cl:
            continue
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
        add_run(p, titulo_cl, bold=True)
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
        add_run(p, corpo_cl)

    # --- CLÁUSULA X (via DB) ---
    titulo_x, corpo_x   = _get_cl("x_corpo", is_unipessoal)
    titulo_p1, corpo_p1 = _get_cl("x_p1", is_unipessoal)
    titulo_p2, corpo_p2 = _get_cl("x_p2", is_unipessoal)

    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
    add_run(p, titulo_x, bold=True)

    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
    add_run(p, corpo_x)

    if corpo_p1:
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
        add_run(p, titulo_p1 + ":", bold=True)
        add_run(p, " " + corpo_p1)

    if corpo_p2:
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
        add_run(p, titulo_p2 + ":", bold=True)
        add_run(p, " " + corpo_p2)

    # --- CLÁUSULAS XI a FORO ---
    cidade_foro     = title_case(end_comercial.get("cidade", "Francisco Beltrão"))
    estado_foro_uf  = end_comercial.get("estado", "PR").strip().upper()
    estado_foro     = ESTADOS_NOMES.get(estado_foro_uf, estado_foro_uf)

    _ROMANOS = ["I","II","III","IV","V","VI","VII","VIII","IX","X","XI","XII","XIII","XIV","XV",
                "XVI","XVII","XVIII","XIX","XX","XXI","XXII","XXIII","XXIV","XXV"]

    clausulas_extras = dados.get("clausulas_extras", [])
    # Numeração: XI=11, XII=12, XIII=13, XIV(ME)=14 → extras → foro
    _num_apos_xiii = 14 if microempresa else 13   # último número antes dos extras
    _num_foro      = _num_apos_xiii + 1 + len(clausulas_extras)
    num_foro       = _ROMANOS[_num_foro - 1]

    # --- CLÁUSULAS XI, XII, XIII (via DB) ---
    for cod in ("xi", "xii", "xiii"):
        titulo_cl, corpo_cl = _get_cl(cod, is_unipessoal)
        if not titulo_cl:
            continue
        titulo_pu, corpo_pu = _get_cl(cod + "_pu", is_unipessoal)
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
        add_run(p, titulo_cl, bold=True)
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL,
                             space_after=SPC_NORMAL if not corpo_pu else SPC_CAPITAL)
        add_run(p, corpo_cl)
        if corpo_pu:
            p = doc.add_paragraph()
            set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
            add_run(p, titulo_pu, bold=True)
            add_run(p, " - " + corpo_pu)

    # --- CLÁUSULA XIV (ME/EPP) via DB ---
    if microempresa:
        titulo_xiv, corpo_xiv = _get_cl("xiv_me", is_unipessoal)
        if titulo_xiv:
            p = doc.add_paragraph()
            set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
            add_run(p, titulo_xiv, bold=True)
            p = doc.add_paragraph()
            set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
            add_run(p, corpo_xiv)

    # --- Cláusulas extras (vêm ANTES do foro) ---
    for i, extra in enumerate(clausulas_extras):
        num_rom = _ROMANOS[_num_apos_xiii + i]   # XIV+1, XIV+2, ...
        titulo_extra = extra.get("titulo", "")
        corpo_extra  = extra.get("corpo", "")
        if not titulo_extra.upper().startswith("CLÁUSULA"):
            titulo_extra = f"CLÁUSULA {num_rom} - {titulo_extra.upper()}"
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
        add_run(p, titulo_extra, bold=True)
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
        add_run(p, corpo_extra)

    # --- Foro (sempre último) ---
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
    add_run(p, f"CLÁUSULA {num_foro} - DO FORO", bold=True)
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
    add_run(p, f"Fica eleito o Foro da Comarca de {cidade_foro} - {estado_foro}, para qualquer ação "
               f"fundada neste contrato, renunciando-se a qualquer outro por muito especial que seja.")

    # --- Fecho ---
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
    add_run(p, "E por estarem em perfeito acordo, em tudo que neste instrumento particular foi "
               "lavrado, obrigam-se a cumprir o presente ato constitutivo, e assinam o presente "
               "em 02 (duas) vias de igual teor e forma.")

    # --- Data ---
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         left_indent=8890, first_line_indent=0,
                         space_after=SPC_DATA)
    add_run(p, f"{cidade_foro} - {estado_foro}, {data_fmt}.")

    # --- Assinaturas (tabela sem bordas, 2 por linha, centralizado) ---
    def papel_socio(s: dict) -> str:
        gen = inferir_genero(s)
        adm = s.get("administrador", False)
        if gen == "f":
            return "Sócia/Administradora" if adm else "Sócia"
        return "Sócio/Administrador" if adm else "Sócio"

    def bloco_assinatura(cell, socio: dict):
        """Preenche uma célula com traço, nome e papel centralizados, espaçamento zero."""
        cell.text = ""
        # Traço
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after  = 0
        add_run(p, "_____________________________")
        # Nome
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = 0
        p2.paragraph_format.space_after  = 0
        add_run(p2, socio["nome"].upper(), bold=True)
        # Papel
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_before = 0
        p3.paragraph_format.space_after  = 0
        add_run(p3, papel_socio(socio), bold=True)

    # Agrupa sócios em pares (2 por linha)
    LARGURA_TABELA = 9374
    largura_col    = LARGURA_TABELA // 2

    i = 0
    while i < len(socios):
        par = socios[i:i+2]
        ncols = len(par)

        tbl_assin = doc.add_table(rows=1, cols=ncols)
        remover_bordas_tabela(tbl_assin)
        set_table_width(tbl_assin, LARGURA_TABELA)

        for j, socio in enumerate(par):
            cell = tbl_assin.rows[0].cells[j]
            set_col_width(cell, largura_col)
            bloco_assinatura(cell, socio)

        i += 2

    doc.save(caminho_saida)
    print(f"✓ Contrato gerado: {caminho_saida}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python3 gerar_contrato.py <ficha.json> [saida.docx]")
        sys.exit(1)

    json_path = sys.argv[1]
    if not os.path.exists(json_path):
        print(f"Erro: arquivo não encontrado: {json_path}")
        sys.exit(1)

    with open(json_path, "r", encoding="utf-8") as f:
        dados = json.load(f)

    razao = dados.get("empresa", {}).get("razaoSocial", "contrato")
    slug  = "".join(c if c.isalnum() or c in " -" else "" for c in razao)
    slug  = slug.strip().replace(" ", "_")[:50]
    saida = sys.argv[2] if len(sys.argv) >= 3 else f"Contrato_{slug}.docx"

    gerar_contrato(dados, saida)
